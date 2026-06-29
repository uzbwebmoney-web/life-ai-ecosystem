from __future__ import annotations

import io
import logging
import re
from datetime import datetime

from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.entities import User, WorkspaceDocument
from app.services.household_service import effective_data_user_ids

logger = logging.getLogger(__name__)

_CHUNK_SIZE = 800
_MAX_SNIPPET = 600


def extract_text_from_bytes(data: bytes, *, filename: str = "", mime: str = "") -> str:
    name = (filename or "").lower()
    mime_l = (mime or "").lower()
    if name.endswith(".txt") or name.endswith(".md") or mime_l.startswith("text/"):
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")

    if name.endswith(".docx") or "wordprocessingml" in mime_l:
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            logger.warning("DOCX extract failed: %s", exc)
            return ""

    if name.endswith(".pdf") or mime_l == "application/pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            parts = []
            for page in reader.pages[:50]:
                parts.append(page.extract_text() or "")
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("PDF extract failed: %s", exc)
            return ""

    try:
        return data.decode("utf-8", errors="ignore")[:50000]
    except Exception:
        return ""


def _score_chunk(query: str, text: str) -> int:
    q_tokens = {t for t in re.findall(r"\w{3,}", query.lower())}
    if not q_tokens:
        return 0
    lower = text.lower()
    return sum(1 for tok in q_tokens if tok in lower)


def _chunk_text(text: str, size: int = _CHUNK_SIZE) -> list[str]:
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for para in paragraphs:
        if len(buf) + len(para) + 1 <= size:
            buf = f"{buf}\n{para}".strip()
        else:
            if buf:
                chunks.append(buf)
            buf = para[:size]
    if buf:
        chunks.append(buf)
    return chunks


async def ingest_workspace_document(
    session: AsyncSession,
    user: User,
    *,
    title: str,
    filename: str,
    mime_type: str,
    file_bytes: bytes,
    telegram_file_id: str | None = None,
    shared: bool = False,
) -> WorkspaceDocument:
    text = extract_text_from_bytes(file_bytes, filename=filename, mime=mime_type)
    household_id = user.household_id if shared and user.household_id else None
    doc = WorkspaceDocument(
        user_id=user.id,
        household_id=household_id,
        title=title.strip()[:255] or filename[:255] or "Document",
        filename=filename[:255],
        mime_type=(mime_type or "")[:128],
        telegram_file_id=telegram_file_id,
        extracted_text=text[:200_000],
        char_count=len(text),
        updated_at=datetime.utcnow(),
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)
    return doc


async def list_workspace_documents(
    session: AsyncSession,
    user: User,
    *,
    limit: int = 20,
) -> list[WorkspaceDocument]:
    user_ids = await effective_data_user_ids(session, user)
    household_ids: set[int] = set()
    if user.household_id:
        household_ids.add(user.household_id)
    clauses = [WorkspaceDocument.user_id.in_(user_ids)]
    if household_ids:
        clauses.append(WorkspaceDocument.household_id.in_(household_ids))
    rows = (
        await session.execute(
            select(WorkspaceDocument)
            .where(
                WorkspaceDocument.active.is_(True),
                or_(*clauses),
            )
            .order_by(WorkspaceDocument.updated_at.desc())
            .limit(limit)
        )
    ).scalars().all()
    return list(rows)


async def search_workspace(
    session: AsyncSession,
    user: User,
    query: str,
    *,
    limit: int = 5,
) -> list[tuple[WorkspaceDocument, str, int]]:
    pattern = f"%{query.strip()}%"
    user_ids = await effective_data_user_ids(session, user)
    household_ids: set[int] = set()
    if user.household_id:
        household_ids.add(user.household_id)
    clauses = [WorkspaceDocument.user_id.in_(user_ids)]
    if household_ids:
        clauses.append(WorkspaceDocument.household_id.in_(household_ids))
    rows = (
        await session.execute(
            select(WorkspaceDocument)
            .where(
                WorkspaceDocument.active.is_(True),
                or_(*clauses),
                or_(
                    WorkspaceDocument.title.ilike(pattern),
                    WorkspaceDocument.filename.ilike(pattern),
                    WorkspaceDocument.extracted_text.ilike(pattern),
                ),
            )
            .order_by(WorkspaceDocument.updated_at.desc())
            .limit(limit * 3)
        )
    ).scalars().all()
    scored: list[tuple[WorkspaceDocument, str, int]] = []
    for doc in rows:
        best_snippet = ""
        best_score = _score_chunk(query, doc.title + " " + doc.extracted_text[:4000])
        for chunk in _chunk_text(doc.extracted_text):
            score = _score_chunk(query, chunk)
            if score > best_score:
                best_score = score
                best_snippet = chunk[:_MAX_SNIPPET]
        if not best_snippet and query.lower() in doc.extracted_text.lower():
            idx = doc.extracted_text.lower().find(query.lower())
            start = max(0, idx - 120)
            best_snippet = doc.extracted_text[start : start + _MAX_SNIPPET]
            best_score = max(best_score, 1)
        if best_score > 0 or query.lower() in doc.title.lower():
            scored.append((doc, best_snippet or doc.extracted_text[:_MAX_SNIPPET], best_score))
    scored.sort(key=lambda x: -x[2])
    return scored[:limit]


async def build_workspace_rag_context(
    session: AsyncSession,
    user: User,
    query: str,
    *,
    limit: int = 4,
) -> str:
    hits = await search_workspace(session, user, query, limit=limit)
    if not hits:
        return ""
    lines = ["📂 Workspace (ваши файлы):"]
    for doc, snippet, _ in hits:
        lines.append(f"• {doc.title} ({doc.filename})")
        if snippet:
            lines.append(f"  {snippet[:400]}")
    return "\n".join(lines)


def format_workspace_hit(doc: WorkspaceDocument, snippet: str, lang: str = "ru") -> str:
    from app.core.i18n import t

    shared = " 👨‍👩‍👧" if doc.household_id else ""
    preview = snippet.replace("\n", " ")[:120] if snippet else "—"
    return t(lang, "workspace_hit_line", title=doc.title, preview=preview) + shared
