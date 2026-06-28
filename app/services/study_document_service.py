from __future__ import annotations

import io
import os
import re
import time
from typing import Literal

ExportFormat = Literal["pdf", "docx", "txt", "md"]

_LAST_DOCS: dict[int, tuple[float, dict[str, str]]] = {}
_DOC_TTL_SECONDS = 3600

_PDF_HINTS = (
    "pdf",
    "пdf",
    "пдф",
    "в pdf",
    "как pdf",
    "файл pdf",
    "сохрани pdf",
    "сделай pdf",
)
_DOCX_HINTS = (
    "docx",
    "word",
    "doc",
    "ворд",
    "в word",
    "как word",
    "файл word",
    "microsoft word",
)
_TXT_HINTS = ("txt", "текстовый файл", "text file", "plain text")
_MD_HINTS = ("markdown", "md", "маркдаун")

_DISCLAIMER_MARKERS = (
    "PDF-файл я напрямую",
    "PDF faylini to'g'ridan-to'g'ri",
    "I can't attach a PDF",
    "I cannot attach a PDF",
    "не могу прикрепить",
    "прикрепить не могу",
    "напрямую прикрепить",
    "If you want, I can",
    "Если хочешь, я могу",
    "Agar xohlasangiz, men",
)


def detect_export_format(text: str) -> ExportFormat | None:
    low = (text or "").lower()
    if any(h in low for h in _PDF_HINTS) or re.search(r"\bpdf\b", low):
        return "pdf"
    if any(h in low for h in _DOCX_HINTS) or re.search(r"\b(docx|doc)\b", low):
        return "docx"
    if any(h in low for h in _MD_HINTS):
        return "md"
    if any(h in low for h in _TXT_HINTS):
        return "txt"
    return None


def is_format_only_request(text: str) -> bool:
    low = (text or "").lower().strip()
    stripped = low
    for token in (
        "pdf",
        "пdf",
        "пдф",
        "docx",
        "word",
        "doc",
        "ворд",
        "txt",
        "md",
        "markdown",
        "сделай",
        "оформи",
        "подготовь",
        "отправь",
        "пришли",
        "скинь",
        "в ",
        "как ",
        "файл",
        "формат",
        "file",
        "format",
        "please",
        "пожалуйста",
    ):
        stripped = stripped.replace(token, " ")
    stripped = re.sub(r"[^\w\s]", " ", stripped)
    stripped = re.sub(r"\s+", " ", stripped).strip()
    return len(stripped) < 25


def extract_document_title(user_text: str, answer: str) -> str:
    for pattern in (
        r"[«\"]([^»\"]+)[»\"]",
        r"(?:по теме|на тему|topic|mavzu)\s+[«\"']?(.+?)[»\"']?(?:$|[.!?])",
        r"конспект\s+(?:по|на)\s+(?:теме\s+)?[«\"']?(.+?)[»\"']?(?:$|[.!?])",
    ):
        match = re.search(pattern, user_text, flags=re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            if title:
                return title[:120]
    for line in answer.splitlines():
        cleaned = re.sub(r"^\d+[\).\s]+", "", line.strip())
        cleaned = cleaned.strip("-• ")
        if cleaned and len(cleaned) > 3:
            return cleaned[:120]
    return "Конспект"


def strip_export_disclaimers(text: str) -> str:
    result = (text or "").strip()
    for marker in _DISCLAIMER_MARKERS:
        idx = result.find(marker)
        if idx <= 0:
            continue
        before = result[:idx].strip()
        prefix = result[:idx]
        if len(before) < 20 and "\n\n" not in prefix:
            continue
        result = before
        break
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result


def store_last_study_document(user_id: int, *, title: str, body: str, query: str) -> None:
    _LAST_DOCS[user_id] = (
        time.time(),
        {"title": title, "body": body, "query": query},
    )


def get_last_study_document(user_id: int) -> dict[str, str] | None:
    row = _LAST_DOCS.get(user_id)
    if not row:
        return None
    ts, payload = row
    if time.time() - ts > _DOC_TTL_SECONDS:
        _LAST_DOCS.pop(user_id, None)
        return None
    return payload


def _find_unicode_font() -> str | None:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    return next((path for path in candidates if os.path.isfile(path)), None)


def _plain_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        line = re.sub(r"^\d+[\).\s]+", "", line)
        line = line.strip("-• ")
        lines.append(line)
    return lines


def build_document_bytes(title: str, body: str, fmt: ExportFormat) -> tuple[bytes, str]:
    clean_body = strip_export_disclaimers(body)
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title).strip() or "study_notes"

    if fmt == "txt":
        content = f"{title}\n\n{clean_body}".encode("utf-8")
        return content, f"{safe_title}.txt"

    if fmt == "md":
        lines = [f"# {title}", ""]
        for line in _plain_lines(clean_body):
            if not line:
                lines.append("")
            elif re.match(r"^\d+\.", line):
                lines.append(line)
            elif len(line) < 80 and not line.startswith(("-", "•")):
                lines.append(f"## {line}")
            else:
                lines.append(line)
        content = "\n".join(lines).encode("utf-8")
        return content, f"{safe_title}.md"

    if fmt == "docx":
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=0)
        for line in _plain_lines(clean_body):
            if not line:
                doc.add_paragraph("")
            elif len(line) < 80 and not line.startswith(("-", "•")) and not re.match(r"^\d+\.", line):
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), f"{safe_title}.docx"

    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    font_path = _find_unicode_font()
    if font_path:
        pdf.add_font("StudyFont", "", font_path)
        pdf.set_font("StudyFont", size=11)
    else:
        pdf.set_font("Helvetica", size=11)

    pdf.multi_cell(0, 8, title)
    pdf.ln(4)
    for line in _plain_lines(clean_body):
        if not line:
            pdf.ln(3)
            continue
        pdf.multi_cell(0, 7, line)
        pdf.ln(1)
    raw = pdf.output()
    if isinstance(raw, bytearray):
        raw = bytes(raw)
    elif isinstance(raw, str):
        raw = raw.encode("latin-1")
    return raw, f"{safe_title}.pdf"
